import os
import time
from deep_translator import GoogleTranslator
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx2txt import process
import re
from concurrent.futures import ThreadPoolExecutor

print("Đang tiến hành xử lý ...")

def chunk_text(text, max_len):
    chunks = []
    while len(text) > max_len:
        idx = text.rfind('.', 0, max_len)
        if idx == -1:
            idx = max_len
        chunk, text = text[:idx], text[idx:]
        chunks.append(chunk.strip())
    chunks.append(text)
    return chunks

def render(name=""):
    url = f"https://read.84000.co/translation/{name}.html?lang=zh"
    # url = "https://www.lotsawahouse.org/tibetan-masters/jamgon-kongtrul/elixir-great-bliss-guru-dewa-chenpo"
    response = requests.get(url)
    soup = BeautifulSoup(response.text, "html.parser")

    # Lấy text từ nội dung web
    text = soup.get_text()

    chunks = chunk_text(text, 5000)  # Chia văn bản thành các phần nhỏ tối đa 5000 ký tự

    # Dịch từng phần và nối lại thành văn bản hoàn thiện
    translated_text = ''
    for chunk in chunks:
        translated_chunk = GoogleTranslator(source='auto', target='vi').translate(chunk)
        # Thêm vào văn bản đã dịch
        translated_text += translated_chunk + ' '

    # Tạo đường dẫn cho thư mục output
    output_folder = "output"
    if not os.path.exists(output_folder):  # Kiểm tra xem thư mục đã tồn tại chưa
        os.makedirs(output_folder)  # Nếu chưa, tạo thư mục
    
    # Tạo đường dẫn cho file output trong thư mục output
    output_path = os.path.join(output_folder, f"{name}.docx")

    # Tạo một file Word mới và lưu nội dung đã dịch vào đó
    doc = Document()
    doc.add_paragraph(translated_text)
    doc.save(output_path)

    # Đọc lại file .docx, xóa ký tự xuống dòng dư thừa và ghi lại nội dung đã chỉnh sửa vào file
    text = process(output_path)
    text = re.sub(r'\n{2,}', '\n', text)
    
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)

# Lấy nội dung web
with ThreadPoolExecutor(max_workers=10) as executor:
  executor.submit(render, "toh846a")

# for i in range(846, 847):
#     print(i)
#     with ThreadPoolExecutor(max_workers=10) as executor:
#         #executor.submit(render, f"toh{i}")
#         executor.submit(render,"elixir-great-bliss-guru-dewa-chenpo")

#         time.sleep(3)

print("Hoàn thành") 