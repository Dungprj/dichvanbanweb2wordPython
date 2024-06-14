import os
import re
import asyncio
import aiohttp
from aiohttp import ClientSession
from aiogoogletrans import Translator
from bs4 import BeautifulSoup
from docx import Document

def chunk_text(text, max_len):
    chunks = []
    while len(text) > max_len:
        idx = text.rfind('.', 0, max_len)
        if idx == -1:
            idx = max_len
        chunk, text = text[:idx], text[idx:]
        chunks.append(chunk.strip())
    chunks.append(text)
    return chunks

async def fetch(session, url):
    async with session.get(url) as response:
        return await response.text()

async def render(urlInput, name, ngonngu = 'vi',folder='output'):
    print(f"Đang tải {name}")
    url = urlInput
    async with ClientSession() as session:
        response_text = await fetch(session, url)
    soup = BeautifulSoup(response_text, "html.parser")

    # Lấy text từ nội dung web
    text = soup.get_text()

    chunks = chunk_text(text, 5000)  # Chia văn bản thành các phần nhỏ tối đa 5000 ký tự

    # Dịch từng phần và nối lại thành văn bản hoàn thiện
    translator = Translator()
    translated_text = ''
    for chunk in chunks:
        translated_chunk = await translator.translate(chunk, src='auto', dest=ngonngu)
        await asyncio.sleep(5)
        translated_text += translated_chunk.text + ' '

    # Tạo đường dẫn cho thư mục output
    output_folder = folder
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Tạo đường dẫn cho file output trong thư mục output
    output_path = os.path.join(output_folder, f"{name}.docx")

    # Tạo một file Word mới và lưu nội dung đã dịch vào đó
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(translated_text)
    run.font.name = 'Times New Roman'
    run.font.size = run.font.size
    doc.save(output_path)

    # Đọc lại file .docx, xóa ký tự xuống dòng dư thừa và ghi lại nội dung đã chỉnh sửa vào file
    doc = Document(output_path)
    paragraphs = doc.paragraphs
    text = '\n'.join([p.text for p in paragraphs])
    text = re.sub(r'\n{2,}', '\n', text)
    
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = run.font.size
    doc.save(output_path)

# print("Đang tiến hành xử lý ...")

# # Chạy hàm bất đồng bộ
# url = 'https://read.84000.co/translation/toh847.html'
# name = 'output_file'
# folder = 'output'
# asyncio.run(render(url, name, folder))
